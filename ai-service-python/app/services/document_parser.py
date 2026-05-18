from __future__ import annotations

import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document as WordDocument
from PIL import Image
from PIL import ImageOps
from pypdf import PdfReader


@dataclass
class ParsedDocument:
    path: Path
    source_name: str
    file_type: str
    content: str


class DocumentParser:
    supported_suffixes = {".pdf", ".docx", ".doc", ".txt", ".md", ".jpg", ".jpeg", ".png"}

    def parse_directory(self, directory: Path) -> list[ParsedDocument]:
        documents: list[ParsedDocument] = []
        if not directory.exists():
            return documents

        for path in sorted(directory.iterdir()):
            if path.is_file() and path.suffix.lower() in self.supported_suffixes:
                documents.append(self.parse_file(path))
        return documents

    def parse_file(self, path: Path) -> ParsedDocument:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            content = self._parse_pdf(path)
        elif suffix == ".docx":
            content = self._parse_docx(path)
        elif suffix == ".doc":
            content = self._parse_doc(path)
        elif suffix in {".jpg", ".jpeg", ".png"}:
            content = self._parse_image(path)
        else:
            content = path.read_text(encoding="utf-8")

        return ParsedDocument(
            path=path,
            source_name=path.name,
            file_type=suffix.lstrip("."),
            content=self._normalize_text(content),
        )

    def _parse_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    def _parse_docx(self, path: Path) -> str:
        document = WordDocument(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())

    def _parse_doc(self, path: Path) -> str:
        worker = Path(__file__).with_name("doc_converter_worker.py")
        if worker.exists():
            try:
                completed = subprocess.run(
                    [sys.executable, "-X", "utf8", str(worker), str(path)],
                    capture_output=True,
                    check=False,
                    encoding="utf-8",
                    errors="replace",
                    timeout=25,
                )
                if completed.returncode == 0 and completed.stdout.strip():
                    return completed.stdout
                if completed.stderr.strip():
                    return (
                        f"旧版 Word 文档 {path.name} 解析失败。"
                        f"请将该文件另存为 DOCX 后再导入知识库。错误信息: {completed.stderr.strip()}"
                    )
            except subprocess.TimeoutExpired:
                return (
                    f"旧版 Word 文档 {path.name} 解析超时。"
                    "请将该文件另存为 DOCX 后再导入知识库。"
                )
            except Exception as exc:
                return (
                    f"旧版 Word 文档 {path.name} 解析失败。"
                    f"请将该文件另存为 DOCX 后再导入知识库。错误信息: {exc}"
                )

        try:
            return self._parse_doc_with_word(path)
        except Exception as exc:
            return (
                f"旧版 Word 文档 {path.name} 暂未解析出正文。"
                f"请确认本机已安装 Microsoft Word 或 WPS，或将该文件另存为 DOCX。"
                f"错误信息: {exc}"
            )

    def _parse_doc_with_word(self, path: Path) -> str:
        import win32com.client

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_docx = Path(temp_dir) / f"{path.stem}.docx"
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            try:
                document = word.Documents.Open(str(path), False, True)
                document.SaveAs(str(temp_docx), FileFormat=16)
                document.Close(False)
            finally:
                word.Quit()

            return self._parse_docx(temp_docx)

    def _parse_image(self, path: Path) -> str:
        tesseract = self._find_tesseract()
        if tesseract is None:
            return (
                f"图片材料 {path.name} 已被识别为知识库附件，但当前系统未安装 OCR 引擎 Tesseract，"
                "暂不能提取图片中的文字。后续可安装 Tesseract 后自动启用图片文字识别。"
            )

        try:
            import pytesseract

            pytesseract.pytesseract.tesseract_cmd = tesseract
            with Image.open(path) as image:
                prepared = self._prepare_image_for_ocr(image)
                return pytesseract.image_to_string(prepared, lang="chi_sim+eng")
        except Exception as exc:
            return f"图片材料 {path.name} OCR 解析失败，错误信息: {exc}"

    def _prepare_image_for_ocr(self, image: Image.Image) -> Image.Image:
        grayscale = ImageOps.grayscale(image)
        enhanced = ImageOps.autocontrast(grayscale)
        width, height = enhanced.size
        return enhanced.resize((width * 2, height * 2))

    def _find_tesseract(self) -> str | None:
        command = shutil.which("tesseract")
        if command:
            return command

        candidates = [
            Path("C:/Program Files/Tesseract-OCR/tesseract.exe"),
            Path("C:/Program Files (x86)/Tesseract-OCR/tesseract.exe"),
        ]
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)
        return None

    def _normalize_text(self, text: str) -> str:
        cleaned_lines = [line.strip() for line in text.splitlines()]
        return "\n".join(line for line in cleaned_lines if line)
