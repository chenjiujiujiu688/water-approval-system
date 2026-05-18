from __future__ import annotations

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SEED_FILE = PROJECT_ROOT / "knowledge-base" / "seeds" / "approval_knowledge_source.md"
RAW_DIR = PROJECT_ROOT / "knowledge-base" / "raw"
DOCX_PATH = RAW_DIR / "water_approval_handbook.docx"
PDF_PATH = RAW_DIR / "water_approval_checklist.pdf"


def load_seed_text() -> str:
    return SEED_FILE.read_text(encoding="utf-8")


def build_docx(content: str) -> None:
    document = Document()
    document.add_heading("涉水审批知识库示例文档", level=1)
    for line in content.splitlines():
        if not line.strip():
            continue
        document.add_paragraph(line.strip())
    document.save(DOCX_PATH)


def build_pdf(content: str) -> None:
    simsun = "C:/Windows/Fonts/simsun.ttc"
    if Path(simsun).exists():
        pdfmetrics.registerFont(TTFont("SimSun", simsun))
        font_name = "SimSun"
    else:
        font_name = "Helvetica"

    pdf = canvas.Canvas(str(PDF_PATH), pagesize=A4)
    width, height = A4
    pdf.setFont(font_name, 12)

    y = height - 50
    for line in content.splitlines():
        text = line.strip()
        if not text:
            y -= 18
            continue
        if y < 50:
            pdf.showPage()
            pdf.setFont(font_name, 12)
            y = height - 50
        pdf.drawString(40, y, text)
        y -= 20

    pdf.save()


def main() -> None:
    RAW_DIR.mkdir(parents=True, exist_ok=True)
    content = load_seed_text()
    build_docx(content)
    build_pdf(content)
    print(f"Generated: {DOCX_PATH}")
    print(f"Generated: {PDF_PATH}")


if __name__ == "__main__":
    main()
